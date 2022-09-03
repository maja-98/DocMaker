import pathlib
import docx
import os
from docx.shared import Inches
import json
import time
from doc_reader import docReader
import sys

with open('config.json', 'r') as f:
  data = json.load(f)

class FileExists(Exception):
    def __init__(self,file_name):
        self.file_name = file_name
    def __str__(self):
     return "File "+self.file_name+ " already there"
class TimeOutError(Exception):
    def __str__(self):
     return "Time Out, Please check your Query"
def queryPrinter(query):
    print("-"*75)
    max_num,max_num_const =75,75
    align_query = ''
    word = ''
    i=0
    line = 1
    while i < len(query):
        letter = query[i]
        if letter == ' ':
            if len(align_query) + len(word) > max_num:
                align_query += '\n'+ word+letter
                line+=1
                max_num = max_num_const * line
            else:
                align_query += word+letter
            i+=1
            word = ''
        elif letter == '[':
            j=i
            start_time = time.time()
            try:
                while query[j]!= ']':

                    letter = query[j]
                    
                    word += letter
                    j+=1
                    if time.time()- start_time> 30:
                        try:
                            raise TimeOutError
                        except Exception as e:
                            print(e)
                word += query[j]
                i=j+1
            except IndexError:
                    print('Please verify Parenthesis are correct in your query')
                    break
        else:
            word+=query[i]
            i+=1
    if len(align_query) + len(word) > max_num:
        align_query += '\n'+ word
    else:
        align_query += word
    print(align_query)
    print("-"*75)

def queryMaker(table,tables):
    if table in tables:
        queryPrinter(tables[table])
    
#Update required: If we capture wrong screenshot, there should be a way to delete that
def docMaker(heading_list):
    try:
        
        print(' Starting docMaker function '.center(75,'*'))
        screenshots_path= data["config"]["screenshot_path"]
        clear_directory= data["config"]["clear_directory"]

        if clear_directory:
            print('-'*50)
            print('WARNING!!!!'.center(50,'-'))
            print('-'*50)
            print('clear directory is True')
            print('your all existing screenshots will deleted...')
            print('Path: '+screenshots_path)
            input('Press any key to continue...')
            print('-'*50)
            
        if not (pathlib.Path(screenshots_path).exists()):
            screenshots_path=input('please configure the correct screenshot saving file path: ')
        if clear_directory:
            print('Clearing directory...')
            for file in os.listdir(screenshots_path):
                os.remove(screenshots_path.strip('/')+'/'+file)
            print('Directory cleared...')            

        main_head = heading_list[0]
        heading_list = heading_list[1:]
        file_name = main_head+'.docx'
        if os.path.exists(file_name):
            try:
                raise FileExists(file_name)
            except Exception as e:
                print(e)
                main_head = input('Enter a new document name: ')
                file_name= main_head.lower().strip('.docx')+'.docx'
        print('-'*50)
        print(('Generating '+ file_name).center(50,'-'))       
        doc=docx.Document()
        doc.core_properties.author= data["config"]["author"]
        a1=doc.add_heading(main_head,level=2)
        a1.alignment = 1
        tables = data["queries"]
        for head in heading_list:            
            doc.add_paragraph(head)
            queryMaker(head,tables)            
            if head:
                value = head
                count = 1
            else:
                count += 1                
            print('Capture screenshot '+str(count)+' for '+value+"...")
            initial_items=set(os.listdir(screenshots_path))            
            while True:
                final_items=set(os.listdir(screenshots_path))
                imgUrl=list(final_items-initial_items)
                if imgUrl :
                    try:
                        doc.add_picture(screenshots_path.strip('/')+'/'+imgUrl[0],width=Inches(7),height=Inches(4))
                        print('screenshot addedd successfully...')
                        break
                    except:
                        print('retrying screenshot add...')
    except KeyboardInterrupt:
        print('File not completed')
        time.sleep(1)
                        
    except Exception as e:
        print(e)
        print('File not completed. Unknow Error')
        time.sleep(1)
    try:
        doc.save(file_name)
        print(("Document "+file_name+" saved").center(50,'-'))
        print('-'*50)
        
    except:
        print("File Not Saved")
        time.sleep(1)
    print(' DocMaker function executed successfully '.center(75,'*'))

documents = data["document"]
headings = [[val]+documents[val] for val in documents ]
for i in range(len(headings)):
    try:
        if headings[i][-1] == "docReader":
            if os.path.exists(headings[i][0]+'.docx'):
                headings[i] = docReader(headings[i][0]+'.docx')
            else:
                print("No such Existing document")
                retry_doc_name = input("Enter correct existing doc path(no need to enter extension): ")
                if os.path.exists(retry_doc_name+'.docx'):
                    headings[i] = docReader(headings[i][0]+'.docx')
                else:
                    print("still not able to find existing doc")
                    print("Skipping current document......")
                    time.sleep(2)
                    continue
    except KeyboardInterrupt:
        print(' DocMaker function Terminated '.center(75,'*'))
        sys.exit()
    except Exception as e:
        print(e)
        print(' DocMaker function Terminated. Unknown Error'.center(75,'*'))
    docMaker(headings[i])
    

